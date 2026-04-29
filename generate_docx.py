import docx

def add_issue(doc, title, summary, scope, deliverables, references, criteria, related):
    doc.add_heading(title, level=1)
    
    doc.add_paragraph("1. Problem Summary: ", style='List Number').add_run(summary)
    doc.add_paragraph("2. Scope: ", style='List Number').add_run(scope)
    
    doc.add_paragraph("3. Deliverables:", style='List Number')
    for d in deliverables:
        doc.add_paragraph(d, style='List Bullet')
        
    doc.add_paragraph("4. References:", style='List Number')
    for r in references:
        doc.add_paragraph(r, style='List Bullet')
        
    doc.add_paragraph("5. Acceptance Criteria:", style='List Number')
    for c in criteria:
        doc.add_paragraph(c, style='List Bullet')
        
    doc.add_paragraph("6. Related Issues: ", style='List Number').add_run(related)
    doc.add_paragraph("") # Space between issues

def main():
    doc = docx.Document()
    
    add_issue(doc, "Issue #1: Test Suite - Cron Scheduler & Manual Sync API Contract",
              "Establish test contracts for the background Cron service that initiates daily synchronization and the manual trigger endpoint. It must be ensured that the system correctly retrieves the active sprint context and validates the HTTP status codes returned after triggering.",
              "QA (Backend Unit & Integration Tests).",
              ["Cron schedule test script: Verify that the cron is triggered at the specified time using mock timestamps.", "API Endpoint Tests: Write test blocks simulating success, error, and not found scenarios for POST /api/v1/scrum-sync/trigger and GET /api/v1/schedules/active-sprint endpoints."],
              ["DFD Sub-processes: 6.1 (Trigger Sync)"],
              ["Test that GET /api/v1/schedules/active-sprint returns 404 Not Found if there is no active sprint in the system.", "Verify that when POST /api/v1/scrum-sync/trigger runs successfully, it initiates the asynchronous process and immediately returns 200 OK (or 202 Accepted).", "Test that the Cron service does not create deadlocks in the system with incorrect time formats or duplicate runs."],
              "#7 Backend - Cron Service & Manual Sync Trigger, #8 Backend - Active Sprint Context Resolver")

    add_issue(doc, "Issue #2: Test Suite - JIRA API Authentication & JQL Query Validation",
              "Test the decryption of team credentials for JIRA integration, authenticate against the external JIRA API, and validate the accuracy of JQL query responses. External dependencies must be isolated by using mock servers.",
              "QA (Backend Integration Tests).",
              ["Authentication tests: Use the data returned from the GET /api/v1/integrations/{teamId}/credentials endpoint in mock JIRA requests.", "JQL Parser tests: Validate mock JIRA paginated JSON responses for POST /api/v1/jira-metrics/issue-query."],
              ["DFD Sub-processes: 6.1, 6.2 (JIRA Metrics)"],
              ["Test that the GET /credentials endpoint returns 403 Forbidden when an incorrect or expired team token is used.", "Assert that paginated data containing 100+ records returned from the mock JIRA API is completely transferred to internal lists.", "Verify that the system does not crash and throws a meaningful exception in case the external API times out."],
              "#9 Backend - Integration Credentials Decryptor, #10 Backend - JIRA JQL Connection Engine")

    add_issue(doc, "Issue #3: Test Suite - GitHub Webhook Security & PR Math Logic",
              "Test HMAC signature validations of Webhook payloads coming from GitHub and verify detailed mathematical calculations of the Story Point performance engine with 100% coverage.",
              "QA (Backend Unit Tests & Security Tests).",
              ["Webhook Security Tests: Send requests to the POST /api/v1/github/webhook/pr-data endpoint with fake (invalid HMAC) and valid signatures.", "PR Math Engine Tests: Test POST /api/v1/story-points/validate and /merge-status algorithms with boundary values."],
              ["DFD Sub-processes: 6.4 (PR Verification), 6.5 (Story Point Validation)"],
              ["Prove that all webhook requests with incorrect or missing X-Hub-Signature-256 headers are immediately rejected with 401 Unauthorized.", "Test that a ZeroDivisionError is not thrown when the divided (target) number in the student SP calculation is 0, and ensure the denominator is safely managed.", "Assert that the system successfully applies the cap limit (1.0 or 100%) when the formula result (e.g., 15/10 SP) is 1.5."],
              "#13 Backend - GitHub PR Webhook Listener, #14 Backend - PR Verification & Math Engine")

    add_issue(doc, "Issue #4: Test Suite - DB Transactions & Audit Logging Resilience",
              "Test Rollback scenarios in case of errors during bulk writing of sprint data to the database. Guarantee that the logging system operates asynchronously (non-blocking) without blocking the main workflow.",
              "QA (Database Resilience & Integration Tests).",
              ["Bulk Update Transaction Tests: Send data sets intentionally containing constraint violations to the PUT /api/v1/sprint-data/bulk-update method.", "Audit Logging Load Tests: Verify the asynchronous behavior of the POST /api/v1/logs/audit-event endpoint."],
              ["DFD Sub-processes: 6.6 (Prepare Sync & Save)"],
              ["Test that in a bulk update operation containing 50 successful and 1 faulty row (e.g., a field that must not be null is null), the 50 successful records are NOT WRITTEN to the database due to a Rollback.", "Prove that the response time of the main thread does not increase due to logging I/O operations when POST /logs/audit-event is called."],
              "#15 Backend - Sprint Data DB Transactions, #16 Backend - Asynchronous System Audit Logger")

    add_issue(doc, "Issue #5: E2E Automation - UI Integration Settings & Toasts",
              "End-to-end (E2E) testing through the browser for UI forms where users enter GitHub PAT and JIRA Space URL information, form validation processes, and Toast notifications that appear after successful/failed operations.",
              "QA (Frontend E2E Automation - Cypress/Playwright).",
              ["Settings Form E2E Test Script: Input mock data into form elements, enforce validation rules, and submit.", "Global Toast E2E Test Script: Capture Toast components on the DOM according to mock successful/failed API responses."],
              ["UI Components: Settings Form, Toast Provider"],
              ["Assert via DOM manipulation that the type of the PAT input element in the form is 'password' (hidden).", "Identify that an HTML or Zod validation warning appears when the JIRA Space URL field is left empty and the submit button is clicked.", "Assert that a green Toast notification with a 'Success' class appears in the corner of the screen after submitting with valid data (with a mock 200 response)."],
              "#17 Frontend - Integration Settings Form, #20 Frontend - Global Toast Notification System")

    add_issue(doc, "Issue #6: E2E Automation - Performance Dashboards & Chart Rendering",
              "End-to-end verification that tables and charts (Dashboards) containing sprint metrics are properly rendered on the screen after receiving data from the API; along with testing table pagination and column-based sorting functionalities.",
              "QA (Frontend E2E Automation - Cypress/Playwright).",
              ["Leaderboard Table & Sorting E2E Script: Click on table headers and verify the row order in the DOM.", "Chart Rendering E2E Script: Test whether the charting library (Recharts/Chart.js) successfully creates a canvas in the DOM."],
              ["UI Components: Performance Table, Analytics Dashboards"],
              ["Verify that the tbody rows in the DOM update in descending/ascending order when the script clicks on 'Ratio' or 'Story Point' column headers.", "Test that old rows disappear and new mock rows arrive when the 'Page 2' button in the Pagination area is clicked.", "Assert that the <canvas> or <svg> tag is present without errors within the visual analysis (Chart) box."],
              "#21 Frontend - System Audit Logs Table, #22 Frontend - Student Performance Leaderboard, #23 Frontend - Sprint Analytics Charts")

    add_issue(doc, "Issue #7: Backend - Cron Service & Manual Sync Trigger",
              "Develop a Cron service to automatically pull data from JIRA and GitHub at a specific time every day, and a manual trigger endpoint allowing coordinators to trigger the sync manually from the menu whenever needed.",
              "Backend (API & Scheduler).",
              ["Endpoint: POST /api/v1/scrum-sync/trigger → Callable by Cron or the user. Starts the synchronization pipeline.", "Background Service: A daily task service configured using Spring Boot @Scheduled or Node-cron."],
              ["DFD Sub-processes: 6.1"],
              ["The scheduled task (Cron job) must trigger seamlessly at the time specified in application.yml or environment variables.", "Calling POST /scrum-sync/trigger must prompt the system to search for the active sprint context and immediately return 200 OK without keeping the user waiting (Asynchronous workflow).", "It must pass the unit tests covered in Phase 1 Test #1 (Coverage > 80%)."],
              "#1 Test Suite - Cron Scheduler API, #8 Backend - Active Sprint Context Resolver")

    add_issue(doc, "Issue #8: Backend - Active Sprint Context Resolver",
              "When Cron or manual operation is triggered, the system needs to know which sprint (in which date range) to process. Write a service to find the currently ongoing active sprint from the timeline in the D10 database (Schedules).",
              "Backend (Database Query & API).",
              ["Endpoint: GET /api/v1/schedules/active-sprint → Queries the exact today's date between sprint start and end dates and returns the matching Sprint DTO."],
              ["Data Stores: D10 (Schedules)", "DFD Sub-processes: 6.1"],
              ["Using the current valid date, only one sprint with 'Active' status or a matching date range should be found in the database.", "If an active sprint is not found, the system must not perform synchronization, and the API must throw a clear 404 Not Found error."],
              "#1 Test Suite - Cron Scheduler API, #7 Backend - Cron Service")

    add_issue(doc, "Issue #9: Backend - Integration Credentials Decryptor",
              "Before making requests to external systems (JIRA/GitHub), securely extract the encrypted credentials (Token, PAT) belonging to the relevant team (Group/Team) from the D3 (Groups) repository, decrypt them securely using algorithms like AES, and provide them to internal services. Must NOT be publicly accessible.",
              "Backend (Security & Encryption).",
              ["Endpoint: GET /api/v1/integrations/{teamId}/credentials → Takes the corresponding group ID, finds encrypted records, decrypts them, and exposes them for internal use.", "Security Utility: String Encryption/Decryption service."],
              ["Data Stores: D3 (Groups/Credentials)", "DFD Sub-processes: 6.1"],
              ["PAT and Token values must strictly not be stored in plain text in the database (D3).", "If an invalid teamId or one that the system is not authorized to access is requested, the endpoint must return 403 Forbidden or 401 Unauthorized."],
              "#2 Test Suite - JIRA API Authentication, #10 Backend - JIRA JQL Connection Engine")

    add_issue(doc, "Issue #10: Backend - JIRA JQL Connection Engine",
              "Establish the initial connection (Initialize) to the JIRA system using decoded credentials and fetch all tasks of the active sprint from the JIRA API using a JQL query (Jira Query Language).",
              "Backend (External API Integration).",
              ["Endpoint: POST /api/v1/jira-metrics/initialize → Tests the connection and initializes a session.", "Endpoint: POST /api/v1/jira-metrics/issue-query → Sends JQL to the external JIRA API (e.g., /rest/api/2/search) to retrieve issue keys (PROJ-1, PROJ-2) belonging to that sprint page by page."],
              ["DFD Sub-processes: 6.1 → 6.2"],
              ["JIRA API authorization (Basic Auth/Bearer) must be successfully added to HTTP request headers.", "If the result exceeds 100 records, the request must continue to be sent using the paginated cursor logic until all pages (all tasks) are exhausted.", "Any 500 errors occurring on the external API side should be logged, and the process should be safely halted."],
              "#2 Test Suite - JIRA API Authentication, #9 Backend - Integration Credentials Decryptor, #11 Backend - JIRA Payload Parser")

    add_issue(doc, "Issue #11: Backend - JIRA Payload Parser & Metrics Appender",
              "Clean and parse the large and complex JSON data returned from the JIRA API; transform only the data relevant to the project (Assignee, Resolution, Story Points) into internal DTO objects and enrich task details.",
              "Backend (Data Mapping).",
              ["Endpoint: POST /api/v1/jira-metrics/callback → Receives raw JSON data, and returns an array of cleansed objects.", "Endpoint: POST /api/v1/jira-metrics/issue-details → Fetches extra details such as SP and Assignee from JIRA for the parsed base issue keys and merges them."],
              ["DFD Sub-processes: 6.2"],
              ["Hundreds of unnecessary meta-data fields in the JIRA JSON structure must be stripped off; only the fields the system needs should be mapped.", "If nobody is yet assigned to a JIRA task (Assignee = null) or no Story Point is entered, the system must not crash, and these values must be assigned as defaults (e.g., 0 or Empty)."],
              "#10 Backend - JIRA JQL Connection Engine, #12 Backend - GitHub Branch Discovery Engine")

    add_issue(doc, "Issue #12: Backend - GitHub Branch Discovery Engine",
              "Use the fetched work keys from JIRA (e.g., PROJ-123) to perform a matching process with the branch names in the GitHub repository to which the team is connected. Automatically identifies which branch a student is working on.",
              "Backend (GitHub API & Algorithms).",
              ["Endpoint: POST /api/v1/github-discovery/start → Initiates the discovery workflow.", "Endpoint: GET /api/v1/github/branch-query → Pulls all branch names from the repo via GitHub API and finds matches with JIRA IDs (e.g., feature/PROJ-123-login) using Regex."],
              ["DFD Sub-processes: 6.2 → 6.3"],
              ["Branch names containing JIRA IDs (PROJ-123) must be found using Regex (or Substring) logic with a minimized margin of error (e.g., case-insensitive search).", "Matching branch information should be gathered in a memory/data structure to be matched with the respective student (Assignee) information."],
              "#11 Backend - JIRA Payload Parser, #13 Backend - GitHub Webhook Listener")

    add_issue(doc, "Issue #13: Backend - GitHub PR Webhook Listener & HMAC",
              "Develop and secure a Webhook endpoint that provides instantaneous data (events) to the backend when a Pull Request (PR) is opened or merged in GitHub repositories.",
              "Backend (Webhooks & Security).",
              ["Endpoint: POST /api/v1/github/webhook/pr-data → Triggered by GitHub, receives the payload, verifies its signature, and records the PR event."],
              ["DFD Sub-processes: 6.4 (PR Verification)"],
              ["The X-Hub-Signature-256 value in the header of the incoming POST request must be verified using the project settings' secret key. If it cannot be verified, the request must be explicitly rejected with 401 Unauthorized.", "The PR action type (opened, closed, merged) inside the Webhook payload must be correctly parsed by the system.", "It must completely pass the boundary security tests identified in Phase 1 Test #3."],
              "#3 Test Suite - GitHub Webhook Security, #12 Backend - GitHub Branch Discovery Engine, #14 Backend - PR Verification & Math Engine")

    add_issue(doc, "Issue #14: Backend - PR Verification & Story Point Calculator Engine",
              "Verify whether the corresponding branches and incoming PR webhook data have truly been successfully 'Merged' into the main branch (main/master). Consequently, gather the Story Points of successful tasks, divide by the targeted SP, and calculate the performance ratio.",
              "Backend (Business Logic & Math Engine).",
              ["PR Verification Endpoints: POST /api/v1/pr-verification/verify and POST /api/v1/pr-verification/branch-list.", "SP Math Engine Endpoints: POST /api/v1/story-points/validate and POST /api/v1/story-points/merge-status."],
              ["DFD Sub-processes: 6.4 → 6.5"],
              ["The task's points (Story Point) should be included in the student's performance calculation only and strictly if the PR status is 'Merged'. (Closed but unmerged PRs yield no points).", "The mathematical algorithm (Completed SP / Target SP) must apply a cap limit with a maximum of 1.0 (meaning 100%).", "If the student's target SP is set as 0 (zero), a 'Zero Division' error must be caught, and the ratio logically assigned to 0 (Aligns with Test #3)."],
              "#3 Test Suite - GitHub Webhook Security, #13 Backend - GitHub Webhook Listener, #22 Frontend - Student Performance Leaderboard")

    add_issue(doc, "Issue #15: Backend - Sprint Data DB Transactions & Bulk Updater",
              "Gather validated and calculated sprint performance data encompassing hundreds of rows within a single asynchronous process (prepare) and commit it into the D6/D8 database repositories via a (bulk update) database transaction. Preserve ACID properties against failures.",
              "Backend (Database Management).",
              ["Prep Endpoint: POST /api/v1/records/prepare-sync → Subjects all gathered data to a final validation check and preps it into bulk format.", "Persistence Endpoint: PUT /api/v1/sprint-data/bulk-update → Persists/updates the prepped data into the database inside a transaction block."],
              ["DFD Sub-processes: 6.5 → 6.6", "Data Stores: D6 (Sprint Data)"],
              ["Data transfer must be conducted purely using 'Bulk Insert / Batch Update' techniques to prevent exhausting database performance.", "If a single Foreign Key collision or error happens during the Transaction phase, the complete rollback of the entire bulk operation must be guaranteed.", "It must successfully fulfill the Test #4 scenarios seen in Phase 1."],
              "#4 Test Suite - DB Transactions Resilience, #16 Backend - Asynchronous System Audit Logger")

    add_issue(doc, "Issue #16: Backend - Asynchronous System Audit Logger",
              "Asynchronously write the success/failure state of each individual stage of the synchronization pipeline, external API calls, and the overarching system execution performance to the system log (D9) without blocking the main workflow.",
              "Backend (Logging & Async Processing).",
              ["Endpoint: POST /api/v1/logs/audit-event → Returns instantly by dispatching the error or info message into an asynchronous queue or thread."],
              ["DFD Sub-processes: 6.6", "Data Stores: D9 (System Logs)"],
              ["Triggering the logging service must operate with a 'Fire and Forget' methodology, levying no extra latency (ms overhead) onto the main synchronization cadence.", "Logs persisted to the database (D9) must be classified by Event Type (Error, Info, Warning) and stored in a ready-to-filter format.", "403 and 500 errors stemming from external APIs must be logged as 'Critical Error' complete with a detailed stack trace."],
              "#4 Test Suite - DB Transactions Resilience, #15 Backend - Sprint Data DB Transactions, #21 Frontend - System Audit Logs Table")

    add_issue(doc, "Issue #17: Frontend - Integration Settings Form & Zod Validation",
              "Establish a UI interface allowing Team Leaders or Coordinators to securely enter integration keys such as GitHub PAT and JIRA Space into the system, and execute client-side validations (Zod/Yup, etc.).",
              "Frontend (UI Layout & Form Handling).",
              ["Integration Settings Form UI: Component containing input fields and a submit button.", "Validation Logic: Definition of form input constraints (cannot be blank, max length, etc.).", "API Call: Secure transmission of validated data towards APIs (via PUT requests) upon form submission."],
              ["UI Component Architecture"],
              ["GitHub PAT and JIRA Token input locations must reside on the screen masked (hidden password type).", "When the user attempts to submit the form while leaving mandatory fields blank, the form component must show distinct and lucid error messages right on the UI without consulting the API.", "Upon reaching a successful API integration, the form must clear itself while triggering the Toast component in Issue #20.", "It must successfully accomplish the E2E Test #5 mapped scenarios."],
              "#5 E2E Automation - UI Integration Settings, #20 Frontend - Global Toast Notification System")

    add_issue(doc, "Issue #18: Frontend - Live Integration Status Indicator UI",
              "Serve a visual indicator component (e.g., green/red dot status) to instantly project the active connectivity state of the JIRA and GitHub bindings native to the chosen group on the advisor's dashboard.",
              "Frontend (UI Component & Status API Binding).",
              ["Status Indicator Component: Visual UX component.", "API Data Fetch Hook: React hook/construct bridging GET /api/v1/integrations/status responses directly into the component's state."],
              ["DFD Sub-processes: 6.1 (Status Fetch)", "Backend API: Feeding through the Status API logic per Issue #9."],
              ["The component must manifest a green 'Connected' status upon observing `connected: true` from the backend; conversely, missing integration or decayed tokens should manifest a red 'No Connection / Disconnected' badge.", "The 'Last Synced' timestamp must visually accompany the component format string (e.g., 10 mins ago).", "Pending API checks must cast a visual Skeleton Loading state."],
              "#9 Backend - Integration Credentials Decryptor")

    add_issue(doc, "Issue #19: Frontend - 'Sync Now' Manual Control & Loading States",
              "Provide a manual control button and in-process UI feedbacks for the admin or advisor allowing them to arbitrarily prompt the traditionally background-cron-dependent synchronization framework whenever circumstances demand.",
              "Frontend (UI/UX & State Management).",
              ["Manual 'Sync Now' Button Component.", "API Event Handler: Routine dispatching POST /api/v1/scrum-sync/trigger backend calls corresponding to click events."],
              ["Backend Service: Issue #7 (POST /scrum-sync/trigger)"],
              ["Activating the button must launch a sync request.", "The button must fall back to a 'Disabled' state casting an animated Spinner (loading icon) internally until the underlying API transaction evaluates (or hits Timeout).", "Processing outcomes, whether victorious or unsuccessful, must report to the Issue #20 Global Toast system."],
              "#7 Backend - Cron Service & Manual Sync Trigger, #20 Frontend - Global Toast Notification System")

    add_issue(doc, "Issue #20: Frontend - Global Toast Notification System",
              "Found the bedrock for a Global Toast Notification construct to appear uniformly at the top-right/bottom-right corners reflecting application-wide events like: Sync initiation, completion, successful form retention, or isolated API interruptions.",
              "Frontend (Providers, Context & UI Rendering).",
              ["Standard Toast Provider/Context architectural wrap.", "Success (Green), Error (Red), and Info (Blue) UI display variations."],
              ["Global UI Elements"],
              ["The structural Provider must stand accessible via highly usable, global functional calls like `toast.success('Message')` decoupled strictly from specific page hierarchies.", "Dismissible intervals must ensure notifications self-immolate without explicit user interaction within typically 3-5 seconds.", "Unduly long text descriptors should gracefully word-wrap without shattering HTML containment dimensions.", "Must endure Toast manifestation routines stipulated down into E2E Test #5 mappings."],
              "#5 E2E Automation - UI Integration Settings, #17 Frontend - Integration Settings Form, #19 Frontend - 'Sync Now' Manual Control")

    add_issue(doc, "Issue #21: Frontend - System Audit Logs Table & Pagination",
              "Parade asynchronous system logs compiled off the backend (pulled out of D9 Database constraints) inside a polished and eminently readable Data Grid (Table) available onto the Advisor’s dashboard. Incorporate pagination boundaries and error filtering buttons to circumvent UI data gluts.",
              "Frontend (UI Layout, API Binding, Local State).",
              ["System Logs Table Component (Tailwind/Material UI constructed).", "Filtering Logic: Operational drops/toggles distinguishing Success instances vis-a-vis Errors.", "Pagination controls anchoring `page=2` API data transitions."],
              ["Backend Service: Issue #16 Audit Logger records"],
              ["Table interfaces must delineate hard columns matching: Date (Timestamp), Event Type (Color-Coded), and System Message.", "Selecting the 'Filter by Error' logic toggle must instantly refresh the UI grid listing strictly Warning/Error log instances.", "Panning forward/backward over pagination sequences should actuate promptly without lag refreshing payload spans natively. (Should traverse E2E Test #6 validations)."],
              "#6 E2E Automation - Dashboards, #16 Backend - Asynchronous System Audit Logger")

    add_issue(doc, "Issue #22: Frontend - Student Performance Leaderboard Data Grid",
              "Sublimate heavily formulated mathematical ratios natively drawn off of backend services (i.e. roughly 85%, 12/15 SP) directly into a visually prominent Leaderboard Data Grid to aid an Advisor's evaluation; equip this framework with generic name/ratio sorting matrices for operational fluency.",
              "Frontend (UI Layout & Sorting Algorithms).",
              ["Performance Table Component (Fully Responsive UI via Tailwind).", "Reversible, State-oriented Sorting apparatus.", "API networking hook funneling dynamic student data points over."],
              ["Backend Service: Issue #14 PR Verification & Math Engine"],
              ["Rendered columns allocating Student Name, Target/Assigned SP, Accomplished SP, and generalized Success Ratios (%) must perfectly colligate with API interfaces visually.", "Visual conditions heavily weighting ~100% outputs in green, against <40% failing states depicted in hard red/orange (Conditional Formatting) MUST characterize the UI layout naturally.", "Clicking specific category headers (i.e., 'Ratio') should routinely alternate ascending/descending alignment schemas properly.", "DOM Sorting stipulations noted in E2E Test #6 stand absolutely mandatory to hurdle."],
              "#6 E2E Automation - Dashboards, #14 Backend - PR Verification & Story Point Calculator Engine")

    add_issue(doc, "Issue #23: Frontend - Sprint Analytics Charts Visualization",
              "Depict a group's SP attrition framework (Burn-down) or overall metrics using compelling Data Visualizations (Charts) located on the advising panoramic board rather than parading stark numeral arrays natively.",
              "Frontend (Data Visualization).",
              ["Analytics Chart Component: Employing either Recharts or Chart.js foundations yielding native semantic Bar or localized Pie visual constructs.", "Internal format structuring morphing backend sprint-data metrics right down into generic structural structures (Label/Value arrays)."],
              ["Backend Service: Feeding off Issue #15 (Sprint Data Bulk Updater outputs)"],
              ["Integrated client-chart-libraries MUST functionally assimilate without exacting ruinous performance loads upon routine page navigations.", "Sub-surface raw data inputs must gracefully expose upon mouse-hover mechanics signaling corresponding informational Toolboxes/Tooltips accurately.", "Total metric vacuums derived across an ongoing sprint MUST NOT yield shattered bounds or empty frames, but universally resolve displaying a cleanly noted 'Empty State' conveying 'No data pooled for this sprint yet'."],
              "#6 E2E Automation - Dashboards, #15 Backend - Sprint Data DB Transactions")

    add_issue(doc, "Issue #24: Frontend - Central Error Handling & API Resilience (Bonus/Stabilization)",
              "Preempt unhandled degradations arising natively during outbound/inbound Frontend-API request interfaces; barring instances of explicit unauthorized responses (401), hard restrictions (403), catastrophic failures (500) or missing nodes (404) sinking natively into silence. Frame a robust Central Interceptor resolving comprehensive data loops right back at users via visual queues.",
              "Frontend (Axios/Fetch Configuration & Error State Management).",
              ["Global Request & Response Interceptors configuration parameters bound up inside Axios/Fetch environments.", "Universally triggerable 403/Access Denied UI error page layouts or states."],
              ["Core Architecture"],
              ["Realizing a 401 Unauthorized API error (due to JWT invalidity) MUST trigger interceptive maneuvers compelling forced Client-side redirects routing backward toward Login interfaces natively.", "White screen of death manifestations deriving from cascading backend 500 errors must be forcibly countered natively utilizing internal callbacks initiating generalized Global Toasts (from Issue #20) declaring generic 'Server Failure Occurred' variables securely.", "Anticipating explicit 404 derivations (i.e. targeting GET /active-sprint) must unconditionally trip standard Empty State frameworks signaling localized messages natively portraying 'No Active Sprint Registered'."],
              "#All Frontend API Interfaces, #20 Frontend - Global Toast Notification System")

    doc.save("process6_translated_issues.docx")

if __name__ == "__main__":
    main()
